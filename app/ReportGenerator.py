import base64
import SimpleITK as sitk
from radiomics import featureextractor
import pydicom
import requests
from fhir.resources.observation import Observation
from fhir.resources.reference import Reference
from fhir.resources.codeableconcept import CodeableConcept
from fhir.resources.coding import Coding
from fhir.resources.diagnosticreport import DiagnosticReport
from fhir.resources.binary import Binary
from fhir.resources.quantity import Quantity 
from docx import Document
from decimal import Decimal
import json
import nibabel as nib
import numpy as np


unit_dict = {
    'original_shape_Volume': 'mm³',
    'original_firstorder_Entropy': None,  # No unit for dimensionless features
    'original_firstorder_Kurtosis': None,  # No unit for histogram features
    'original_firstorder_Skewness': None   # No unit for histogram features
}

params = {
    'imageType': {'Original': {}},
    'setting': {
        'binWidth': 25,
        'resampledPixelSpacing': None,
        'force2D': False
    },
    'featureClass': {
        'shape': None,  # Extract all shape features
        'firstorder': None,  # Extract all first-order features
        'glcm': None,  # Extract all GLCM features
        'glrlm': None,  # Extract all GLRLM features
        'glszm': None  # Extract all GLSZM features
    }
}

extractor = featureextractor.RadiomicsFeatureExtractor(params)

def getImageID(url,study_id):
    print("getImageID ...",url + "/" + study_id)
    resp = requests.get(url + "/" + "ImagingStudy/"+study_id)
    imaging = resp.text
    print(imaging)
    imaging = json.loads(imaging)
    series = imaging['series']
    print("series::",series)
    image_id = series[0]
    print("image_id::",image_id["uid"])
    image_id = image_id["uid"]
    return image_id
def nifti_to_sitk(nifti_file_path):
    # Load NIfTI file using nibabel
    nii = nib.load(nifti_file_path)
    data = nii.get_fdata()

    # Convert numpy array to SimpleITK Image
    sitk_image = sitk.GetImageFromArray(np.asanyarray(data))
    sitk_image.SetSpacing(nii.header.get_zooms()[:3])
    sitk_image.SetOrigin(nii.affine[:3, 3])
    return sitk_image

def convert_to_binary_mask(sitk_mask, positive_label):
    """Convert a mask to a binary format where only the specified `positive_label` is 1, others are 0."""
    # Convert the image to a numpy array
    mask_array = sitk.GetArrayFromImage(sitk_mask)
    # Create binary mask
    binary_mask_array = (mask_array == positive_label).astype(int)
    # Convert back to SimpleITK image
    binary_mask_sitk = sitk.GetImageFromArray(binary_mask_array)
    binary_mask_sitk.CopyInformation(sitk_mask)
    return binary_mask_sitk

class DecimalEncoder(json.JSONEncoder):
    def default(self, obj):
        if isinstance(obj, Decimal):
            return float(obj)  # or use str(obj) if precision is critical
        if isinstance(obj, bytes):
            return base64.b64encode(obj).decode('utf-8')  # convert bytes to base64 string
        return super(DecimalEncoder, self).default(obj)

def load_nifti_as_sitk(path):
    """Load a NIfTI file as a SimpleITK Image."""
    return sitk.ReadImage(path)   

def extract_features(image_path, mask_path,label_id):
    sitk_image = load_nifti_as_sitk(image_path)
    sitk_mask = load_nifti_as_sitk(mask_path)

    # Convert the mask to binary
    binary_mask = convert_to_binary_mask(sitk_mask,label_id)

    # nii = nib.load(mask_path)
    # pred_nrrd_lung = nii.get_fdata() 
    extracted_features = extractor.execute(sitk_image, binary_mask)
    # Filter out non-numeric and metadata entries
    features = {k: v for k, v in extracted_features.items() if isinstance(v, (int, float))}
    return features

def create_observation(feature_name, feature_value, patient_id, imaging_study_id, unit_dict):
    # Create the FHIR Observation resource
    unit = unit_dict.get(feature_name, None)
    observation = Observation(
        status='final',
        code=CodeableConcept(
            coding=[Coding(
                system='http://terminology.hl7.org/CodeSystem/observation-category',
                code=feature_name,
                display='Radiomic Feature'
            )],
            text = feature_name
        ),
        subject=Reference(
            reference=f'/{patient_id}'
        ),
        # encounter=Reference(
        #     reference=f'ImagingStudy/{imaging_study_id}'
        # ),
        valueQuantity=Quantity(
            value=float(feature_value),  # Now safely converted to float
            unit=unit,
            system='http://unitsofmeasure.org',
            code='1'
        )
    )
    return observation

# def create_observation(feature_name, feature_value, patient_id, imaging_study_id):
#     observation = Observation(
#         status='final',
#         code=CodeableConcept(
#             coding=[Coding(
#                 system='http://terminology.hl7.org/CodeSystem/observation-category',
#                 code=feature_name,
#                 display='Radiomic Feature'
#             )]
#         ),
#         subject=Reference(
#             reference=f'Patient/{patient_id}'
#         ),
#         valueQuantity=Quantity(
#             value=float(feature_value),
#             unit='Unit',
#             system='http://unitsofmeasure.org',
#             code='1'
#         )
#     )
#     return observation

def generate_report_doc(features):
    doc = Document()
    doc.add_heading('Lung Nodule Analysis Report', 0)
    for key, value in features.items():
        doc.add_paragraph(f'{key}: {value}')
    return doc

# Example usage:
def post_fhir_resource(resource,FHIR_SERVER_URL,headers):
    """Post a FHIR resource to the server and return the resource ID."""
    # headers = {'Content-Type': 'application/fhir+json'}
    resource_json = json.dumps(resource.dict(), cls=DecimalEncoder)
    try:
        response = requests.post(f'{FHIR_SERVER_URL}/{resource.resource_type}', headers=headers, data=resource_json)
        # response = requests.post(f'{FHIR_SERVER_URL}/{resource.resource_type}', headers=headers, json=resource.dict())
        response.raise_for_status()  # Raise an exception for HTTP errors
        print("Response Status Code:", response.status_code)
        print("Response JSON:", response.json())

        # Now check if 'id' is in the response and handle if it's not
        response_data = response.json()
        if 'id' in response_data:
            return response_data['id']
        else:
            raise ValueError("No 'id' found in response, check error details above.")
            return response.json()['id']
    except requests.exceptions.HTTPError as error:
        print("Failed to post resource:", error.response.status_code)
        print("Response:", error.response.text)  # Log the error response
        raise
# FHIR server base URL
def process(FHIR_SERVER_URL,patient_id,imaging_study_id,image_path,mask_path,label_id,inference_findings):

    # FHIR_SERVER_URL = FHIR_SERVER_URL

    # patient_id = '123'
    # imaging_study_id = '456'
    # image_path = 'path_to_dicom_image.dcm'
    # mask_path = 'path_to_nodule_mask.nii'
    image_id = getImageID(FHIR_SERVER_URL,imaging_study_id)
    features = extract_features(image_path, mask_path,label_id)
    observations = [create_observation(name, value, patient_id, imaging_study_id, unit_dict) for name, value in features.items()]
    # print("Observation   ::",observations)
    headers = {'Content-Type': 'application/fhir+json'}
    observation_ids = [post_fhir_resource(obs,FHIR_SERVER_URL,headers) for obs in observations]

    # report_doc = generate_report_doc(features)
    # report_doc_path = 'Lung_Nodule_Report.docx'
    # report_doc.save(report_doc_path)

    # # Post the report document as a binary resource
    # with open(report_doc_path, 'rb') as f:
    #     # binary = Binary(contentType='application/msword', data=f.read().encode('base64'))
    #     binary = Binary(contentType='application/msword', data=f.read())
    #     headers = {'Content-Type': 'application/msword'}
    #     binary_id = post_fhir_resource(binary,FHIR_SERVER_URL,headers)

    # Create and post diagnostic report
    diagnostic_report = DiagnosticReport(
        status='final',
        code=CodeableConcept(
            coding=[Coding(
                system='http://loinc.org',
                code='18748-4',
                display='Radiology Report'
            )]
        ),
        subject=Reference(
            reference=f'/{patient_id}'
        ),
        conclusion=inference_findings,
        conclusionCode = [CodeableConcept(
            coding=[Coding(
                system='http://loinc.org',
                code=image_id,  # This might be incorrect, typically should be a diagnostic code, not an ID
                display=inference_findings
            )]
        )],
        # encounter=Reference(
        #     reference=f'ImagingStudy/{imaging_study_id}'
        # ),
        result=[Reference(reference=f'Observation/{obs_id}') for obs_id in observation_ids]
        # result=[Reference(reference=f'Observation/{obs_id}') for obs_id in observation_ids],
        # presentedForm=[Reference(reference=f'Binary/{binary_id}')]
    )
    print(diagnostic_report)
    headers = {'Content-Type': 'application/fhir+json'}
    diagnostic_report_id = post_fhir_resource(diagnostic_report,FHIR_SERVER_URL,headers)

    print(f'Diagnostic Report ID: {diagnostic_report_id}')
